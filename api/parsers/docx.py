"""DOCX text extraction via python-docx.

Pure CPU; no DB, no network, no agent (iron rules #8, #9, #10 inapplicable).
Iterates the document's paragraphs in order, then table cells in row-major
order, concatenating with `\\n\\n` separators.

Hebrew + bidirectional text is preserved verbatim — python-docx reads the
docx XML directly, so RTL runs and mixed bidi survive without reordering.

Failure modes:
    - Not a zip (corrupt or .doc legacy binary) → `"corrupt"`.
    - Zip is valid but not a docx package → `"corrupt"`.

Image-only DOCX documents — where text lives only inside embedded
pictures — return `""`. The worker handler maps this to
`WorkerErrorClass.ParseEmpty`. Image-PDF/DOCX is M2b #6 OCR territory.

Document structure caveat: this extractor walks the top-level
`document.paragraphs` and `document.tables` collections separately,
losing the interleaving between body paragraphs and tables. For
chunking that's acceptable — chunks operate on token-window slices and
don't depend on document-level paragraph-table interleaving. Headers,
footers, footnotes, comments, and text inside text-boxes are NOT
extracted; M4 #6 may revisit if Priority docs use those surfaces.
"""

from __future__ import annotations

import io
from zipfile import BadZipFile

import docx
from docx.opc.exceptions import PackageNotFoundError

from api.parsers.types import ParserError


def parse_docx(data: bytes) -> str:
    """Extract text from a DOCX byte buffer.

    Returns paragraph text followed by table cell text, joined with
    `\\n\\n`. An empty document returns `""`.

    Raises:
        ParserError("corrupt") — bytes are not a valid DOCX package.
    """
    try:
        document = docx.Document(io.BytesIO(data))
    except (BadZipFile, PackageNotFoundError) as e:
        # BadZipFile fires when the bytes aren't a zip at all (random
        # data, legacy .doc binary, truncated upload). PackageNotFoundError
        # fires when the zip is valid but missing the docx parts (e.g.,
        # a generic .zip uploaded with a .docx extension).
        raise ParserError("corrupt", str(e)) from e

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if text:
                    parts.append(text)

    return "\n\n".join(parts)
